"""Document Renderer — StructuredDocument → PDF / DOCX / Markdown / HTML."""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Optional

from config.paths import get_paths

from .document_model import StructuredDocument


def _safe_filename(title: str, ext: str) -> str:
    base = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip()
    base = base.replace(" ", "_")[:60] or "document"
    return f"{base}_{uuid.uuid4().hex[:8]}.{ext}"


class DocumentRenderer:
    """Platform renderer — no LLM involvement."""

    def render(self, doc: StructuredDocument, format: str = "pdf", output_dir: Optional[Path] = None) -> Path:
        fmt = format.lower().replace(".", "")
        if fmt == "md":
            fmt = "markdown"
        out_dir = output_dir or get_paths().artifacts_dir
        out_dir.mkdir(parents=True, exist_ok=True)

        if fmt == "pdf":
            return self._pdf(doc, out_dir)
        if fmt == "docx":
            return self._docx(doc, out_dir)
        if fmt == "html":
            return self._html(doc, out_dir)
        if fmt in {"markdown", "txt"}:
            return self._markdown(doc, out_dir, ext="md" if fmt == "markdown" else "txt")
        raise ValueError(f"Unsupported document format: {format}")

    def _latin(self, text: str) -> str:
        return (text or "").encode("latin-1", errors="replace").decode("latin-1")

    def _pdf(self, doc: StructuredDocument, out_dir: Path) -> Path:
        from fpdf import FPDF

        path = out_dir / _safe_filename(doc.title, "pdf")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 10, self._latin(doc.title))
        pdf.ln(2)
        if doc.summary:
            pdf.set_font("Helvetica", "I", 11)
            pdf.multi_cell(0, 6, self._latin(doc.summary))
            pdf.ln(4)
        for sec in doc.sections:
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, self._latin(sec.heading))
            pdf.ln(1)
            if sec.body:
                pdf.set_font("Helvetica", size=11)
                for line in sec.body.splitlines() or [" "]:
                    pdf.multi_cell(0, 6, self._latin(line) if line else " ")
            for bullet in sec.bullets:
                pdf.set_font("Helvetica", size=11)
                pdf.multi_cell(0, 6, self._latin(f"- {bullet}"))
            pdf.ln(3)
        pdf.output(str(path))
        if not path.is_file() or path.stat().st_size == 0:
            raise RuntimeError(f"PDF render produced empty file: {path}")
        return path

    def _docx(self, doc: StructuredDocument, out_dir: Path) -> Path:
        from docx import Document

        path = out_dir / _safe_filename(doc.title, "docx")
        d = Document()
        d.add_heading(doc.title, level=1)
        if doc.summary:
            d.add_paragraph(doc.summary)
        for sec in doc.sections:
            d.add_heading(sec.heading, level=min(sec.level + 1, 3))
            if sec.body:
                d.add_paragraph(sec.body)
            for bullet in sec.bullets:
                d.add_paragraph(bullet, style="List Bullet")
        d.save(str(path))
        return path

    def _markdown(self, doc: StructuredDocument, out_dir: Path, ext: str = "md") -> Path:
        path = out_dir / _safe_filename(doc.title, ext)
        parts = [f"# {doc.title}", ""]
        if doc.summary:
            parts.extend([doc.summary, ""])
        for sec in doc.sections:
            level = min(max(sec.level + 1, 2), 4)
            parts.append("#" * level + f" {sec.heading}")
            parts.append("")
            if sec.body:
                parts.append(sec.body)
                parts.append("")
            for b in sec.bullets:
                parts.append(f"- {b}")
            if sec.bullets:
                parts.append("")
        path.write_text("\n".join(parts), encoding="utf-8")
        return path

    def _html(self, doc: StructuredDocument, out_dir: Path) -> Path:
        path = out_dir / _safe_filename(doc.title, "html")

        def esc(s: str) -> str:
            return (
                (s or "")
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )

        parts = [
            "<!DOCTYPE html><html><head><meta charset='utf-8'>",
            f"<title>{esc(doc.title)}</title></head><body>",
            f"<h1>{esc(doc.title)}</h1>",
        ]
        if doc.summary:
            parts.append(f"<p><em>{esc(doc.summary)}</em></p>")
        for sec in doc.sections:
            parts.append(f"<h2>{esc(sec.heading)}</h2>")
            if sec.body:
                for para in sec.body.split("\n\n"):
                    parts.append(f"<p>{esc(para)}</p>")
            if sec.bullets:
                parts.append("<ul>")
                for b in sec.bullets:
                    parts.append(f"<li>{esc(b)}</li>")
                parts.append("</ul>")
        parts.append("</body></html>")
        path.write_text("\n".join(parts), encoding="utf-8")
        return path
