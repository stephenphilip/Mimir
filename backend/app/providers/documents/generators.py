"""Low-level document file generators (no chat/orchestrator coupling)."""

from __future__ import annotations

import csv
import io
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional

from config.paths import get_paths


def _safe_filename(title: str, ext: str) -> str:
    base = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip()
    base = base.replace(" ", "_")[:60] or "document"
    return f"{base}_{uuid.uuid4().hex[:8]}.{ext}"


def generate_pdf(title: str, content: str, output_dir: Optional[Path] = None) -> Path:
    from fpdf import FPDF

    def _safe(text: str) -> str:
        return text.encode("latin-1", errors="replace").decode("latin-1")

    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "pdf")
    path = out_dir / filename

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 10, _safe(title))
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for line in content.splitlines():
        pdf.multi_cell(0, 6, _safe(line) if line else " ")
        pdf.x = pdf.l_margin
    pdf.output(str(path))

    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError(f"PDF generation produced empty file: {path}")
    return path


def generate_markdown(title: str, content: str, output_dir: Optional[Path] = None) -> Path:
    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "md")
    path = out_dir / filename
    body = f"# {title}\n\n{content}\n"
    path.write_text(body, encoding="utf-8")
    return path


def generate_txt(title: str, content: str, output_dir: Optional[Path] = None) -> Path:
    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "txt")
    path = out_dir / filename
    path.write_text(f"{title}\n\n{content}", encoding="utf-8")
    return path


def generate_docx(title: str, content: str, output_dir: Optional[Path] = None) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX generation") from exc

    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "docx")
    path = out_dir / filename

    doc = Document()
    doc.add_heading(title, level=1)
    for para in content.split("\n\n"):
        doc.add_paragraph(para)
    doc.save(str(path))
    return path


def generate_csv(
    title: str,
    rows: List[List[Any]],
    output_dir: Optional[Path] = None,
) -> Path:
    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "csv")
    path = out_dir / filename

    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        for row in rows:
            writer.writerow(row)
    return path


def generate_xlsx(
    title: str,
    rows: List[List[Any]],
    output_dir: Optional[Path] = None,
) -> Path:
    try:
        from openpyxl import Workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX generation") from exc

    out_dir = output_dir or get_paths().artifacts_dir
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(title, "xlsx")
    path = out_dir / filename

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31] or "Sheet1"
    for row in rows:
        ws.append(row)
    wb.save(str(path))
    return path


def parse_tabular_content(content: str) -> List[List[str]]:
    """Best-effort parse of markdown table or line-based CSV in prompt content."""
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    rows: List[List[str]] = []
    for line in lines:
        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= {"-", ":"} for c in cells):
                continue
            rows.append(cells)
        elif "," in line:
            rows.append([c.strip() for c in line.split(",")])
        else:
            rows.append([line])
    if not rows:
        rows = [["Content"], [content[:500]]]
    return rows
